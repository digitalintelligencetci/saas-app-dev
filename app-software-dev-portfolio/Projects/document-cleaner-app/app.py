import streamlit as st
import re
import io
import zipfile
from docx import Document
import PyPDF2
import pandas as pd
from pathlib import Path

# Page configuration
st.set_page_config(
    page_title="Document Cleaner Pro",
    page_icon="🧹",
    layout="wide"
)

def clean_line(line, cleaning_options):
    """Enhanced line cleaning with configurable options"""
    if cleaning_options.get('remove_obsidian_links', True):
        line = re.sub(r'!\[\[.*?\]\]', '', line)
        line = re.sub(r'\[\[(.*?)\]\]', r'\1', line)
    
    if cleaning_options.get('remove_urls', False):
        line = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', line)
    
    if cleaning_options.get('remove_emails', False):
        line = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '', line)
    
    if cleaning_options.get('remove_extra_whitespace', True):
        line = re.sub(r'\s+', ' ', line).strip()
    
    return line

def remove_yaml_frontmatter(lines, remove_yaml=True):
    """Remove YAML frontmatter with option to disable"""
    if not remove_yaml:
        return lines
    
    cleaned = []
    inside_yaml = False
    for line in lines:
        if line.strip() == "---":
            inside_yaml = not inside_yaml
            continue
        if not inside_yaml:
            cleaned.append(line)
    return cleaned

def split_large_text(text, max_bytes=2 * 1024 * 1024, split_by_chunks=False, chunk_size=1000):
    """Enhanced text splitting with multiple options"""
    if split_by_chunks:
        # Split by number of characters
        chunks = []
        for i in range(0, len(text), chunk_size):
            chunks.append(text[i:i + chunk_size])
        return chunks
    else:
        # Original byte-based splitting
        chunks = []
        current = ""
        for line in text.splitlines(keepends=True):
            if len(current.encode('utf-8')) + len(line.encode('utf-8')) > max_bytes:
                chunks.append(current)
                current = ""
            current += line
        if current:
            chunks.append(current)
        return chunks

def process_pdf_file(uploaded_file):
    """Extract text from PDF files"""
    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.splitlines(keepends=True)
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
        return []

def process_file(uploaded_file, cleaning_options, split_options):
    """Enhanced file processing with multiple format support"""
    file_extension = Path(uploaded_file.name).suffix.lower()
    
    # Progress bar
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    status_text.text("Reading file...")
    progress_bar.progress(25)
    
    # Process different file types
    if file_extension == '.docx':
        doc = Document(uploaded_file)
        raw_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        raw_lines = raw_text.splitlines(keepends=True)
        original_text = raw_text
    elif file_extension == '.pdf':
        raw_lines = process_pdf_file(uploaded_file)
        original_text = ''.join(raw_lines)
    else:
        # Handle text files with encoding detection
        try:
            # Try UTF-8 first
            file_content = uploaded_file.read()
            raw_text = file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                # Try other common encodings
                uploaded_file.seek(0)  # Reset file pointer
                raw_text = file_content.decode("latin-1")
            except UnicodeDecodeError:
                uploaded_file.seek(0)  # Reset file pointer
                raw_text = file_content.decode("cp1252")
        
        raw_lines = raw_text.splitlines(keepends=True)
        original_text = raw_text
    
    status_text.text("Cleaning content...")
    progress_bar.progress(50)
    
    # Apply cleaning
    lines = remove_yaml_frontmatter(raw_lines, cleaning_options.get('remove_yaml', True))
    cleaned_lines = [clean_line(line, cleaning_options) for line in lines]
    full_text = ''.join(cleaned_lines)
    
    status_text.text("Splitting content...")
    progress_bar.progress(75)
    
    # Split text
    if split_options.get('split_by_chunks', False):
        chunks = split_large_text(full_text, split_by_chunks=True, chunk_size=split_options.get('chunk_size', 1000))
    else:
        max_bytes = split_options.get('max_bytes', 2 * 1024 * 1024)
        chunks = split_large_text(full_text, max_bytes=max_bytes)
    
    progress_bar.progress(100)
    status_text.text("Processing complete!")
    
    return chunks, full_text, original_text

def structure_paragraphs(text):
    """Structure text into clean paragraphs"""
    # Split by double newlines to identify paragraphs
    paragraphs = re.split(r'\n\s*\n', text)
    
    # Clean and structure each paragraph
    structured_paragraphs = []
    for paragraph in paragraphs:
        # Remove extra whitespace and normalize
        cleaned = re.sub(r'\s+', ' ', paragraph.strip())
        if cleaned:  # Only add non-empty paragraphs
            structured_paragraphs.append(cleaned)
    
    return structured_paragraphs

def create_docx_from_chunks(chunks, filename="cleaned_document.docx"):
    """Create a DOCX file from text chunks"""
    doc = Document()
    
    for i, chunk in enumerate(chunks):
        # Add a heading for each chunk
        if len(chunks) > 1:
            doc.add_heading(f'Section {i+1}', level=1)
        
        # Structure the chunk into paragraphs
        paragraphs = structure_paragraphs(chunk)
        
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text)
        
        # Add a page break between chunks (except for the last one)
        if i < len(chunks) - 1:
            doc.add_page_break()
    
    # Save to bytes buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def display_statistics(original_text, cleaned_text, chunks):
    """Display processing statistics"""
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Original Size", f"{len(original_text):,} chars")
    with col2:
        st.metric("Cleaned Size", f"{len(cleaned_text):,} chars")
    with col3:
        st.metric("Reduction", f"{((len(original_text) - len(cleaned_text)) / len(original_text) * 100):.1f}%")
    with col4:
        st.metric("Chunks Created", len(chunks))

def main():
    st.title("🧹 Document Cleaner Pro")
    st.markdown("Clean and process your documents with advanced options")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Cleaning options
        st.subheader("Cleaning Options")
        cleaning_options = {
            'remove_obsidian_links': st.checkbox("Remove Obsidian links", value=True),
            'remove_urls': st.checkbox("Remove URLs", value=False),
            'remove_emails': st.checkbox("Remove email addresses", value=False),
            'remove_extra_whitespace': st.checkbox("Remove extra whitespace", value=True),
            'remove_yaml': st.checkbox("Remove YAML frontmatter", value=True)
        }
        
        # Split options
        st.subheader("Split Options")
        split_by_chunks = st.checkbox("Split by character count", value=False)
        split_options = {
            'split_by_chunks': split_by_chunks,
            'chunk_size': st.number_input("Characters per chunk", min_value=100, max_value=10000, value=1000, disabled=not split_by_chunks),
            'max_bytes': st.number_input("Max bytes per chunk (MB)", min_value=0.1, max_value=10.0, value=2.0, step=0.1) * 1024 * 1024
        }
        
        # Output format
        st.subheader("Output Format")
        output_format = st.selectbox("Output format", ["docx", "txt", "csv"])
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        uploaded_file = st.file_uploader(
            "Upload your document", 
            type=["docx", "pdf", "txt", "md"],
            help="Supported formats: DOCX, PDF, TXT, MD"
        )
    
    with col2:
        if uploaded_file:
            st.info(f"📄 {uploaded_file.name}")
            st.metric("File Size", f"{uploaded_file.size / 1024:.1f} KB")
    
    if uploaded_file:
        # Process file
        chunks, cleaned_text, original_text = process_file(uploaded_file, cleaning_options, split_options)
        
        # Display statistics
        display_statistics(original_text, cleaned_text, chunks)
        
        # Create output based on format
        if output_format == "docx":
            # Create structured DOCX file
            docx_buffer = create_docx_from_chunks(chunks, f"cleaned_{Path(uploaded_file.name).stem}.docx")
            
            # Download section
            st.success(f"✅ Processed and structured into {len(chunks)} section(s).")
            
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📄 Download Cleaned DOCX",
                    data=docx_buffer,
                    file_name=f"cleaned_{Path(uploaded_file.name).stem}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            # Create ZIP for other formats
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                for i, chunk in enumerate(chunks):
                    filename = f"cleaned_part_{i+1}.{output_format}"
                    zf.writestr(filename, chunk)
            zip_buffer.seek(0)
            
            # Download section
            st.success(f"✅ Processed and split into {len(chunks)} file(s).")
            
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📦 Download Cleaned ZIP",
                    data=zip_buffer,
                    file_name=f"cleaned_output_{Path(uploaded_file.name).stem}.zip",
                    mime="application/zip"
                )
        
        with col2:
            # Preview option
            if st.button("👁️ Preview First Chunk"):
                if chunks:
                    st.text_area("Preview", chunks[0], height=300)
        
        # Export as DataFrame
        if st.checkbox("Export as CSV"):
            df = pd.DataFrame({
                'chunk_number': range(1, len(chunks) + 1),
                'content': chunks,
                'length': [len(chunk) for chunk in chunks]
            })
            csv = df.to_csv(index=False)
            st.download_button(
                label="📊 Download CSV",
                data=csv,
                file_name="cleaned_chunks.csv",
                mime="text/csv"
            )

if __name__ == "__main__":
    main()